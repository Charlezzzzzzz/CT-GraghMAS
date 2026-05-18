"""Format the middle report area of the final application DOCX."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "结题报告" / "附件6：国家级“大学生创新创业训练计划项目”结题申请书_排版版.docx"
OUT = ROOT / "结题报告" / "附件6：国家级“大学生创新创业训练计划项目”结题申请书_排版完成版.docx"

INK = RGBColor(0, 0, 0)
BLUE = RGBColor(31, 77, 120)


TABLE_TITLES = [
    "团队成员分工与合作情况",
    "项目阶段推进与成果形成情况",
    "ADAPT伴学模式中五类智能体的功能分工与算法思维指向",
    "C-O-D-E-R课堂范式的阶段结构与计算思维培养对应关系",
    "知识图谱检索策略测评数据",
    "不同学习者画像下的L2教学智能体路由策略",
    "角色权限与数据流向",
    "CT-GraphMAS三层多智能体提示词设计",
    "ADAPT伴学模式智能体设计",
    "C-O-D-E-R课堂范式智能体设计",
    "项目经费使用与报销汇总",
]

FIGURE_TITLES = [
    "GCCIL2025优秀成果一等奖奖状截图",
    "HWESM 2026论文录用通知截图",
    "EDCS 2026论文录用通知截图",
    "ADAPT伴学协作模式图",
    "ADAPT伴学算法思维培养路径图",
    "C-O-D-E-R范式中的多智能体准主体集群图",
    "C-O-D-E-R课堂范式流程图",
    "CT-GraphMAS三层多智能体系统架构图",
    "知识图谱检索策略实验结果图",
    "不同学习者画像下的多智能体协作热力图",
    "CT-GraphMAS系统业务流程图",
    "默认学号、班级任务、学习画像与入口导航",
    "标准辅导入口与固定滚动对话框",
    "面向课堂流程的C-O-D-E-R与ADAPT适配",
    "L1诊断、L2协同、L3呈现与相关知识",
    "检索、拖动与前置相关节点查看",
    "班级任务分配、课堂交互与任务编排",
    "上传文件、生成相关知识、embedding与图谱视图",
    "系统状态、角色入口与维护概览",
    "学习者画像、多智能体偏好、API与教学模式配置",
    "核心提示词、工作流步骤与运行时设计",
    "管理员图谱视图：Python微观考点、学情节点与依赖关系",
]

MENTOR_COMMENT = (
    "该项目围绕生成式人工智能背景下初中Python编程学习的真实需求展开，选题具有较强的教育实践价值和研究创新性。项目组能够按照立项计划持续推进，围绕ADAPT伴学模式、C-O-D-E-R课堂范式和CT-GraphMAS系统架构形成较为完整的成果链条，并完成知识图谱增强多智能体原型系统开发与初步实验验证。项目实施过程中，团队分工明确，能够主动开展文献梳理、系统设计、数据整理、论文撰写和成果展示，体现出较好的科研训练水平、协作能力和问题意识。相关成果在多智能体教学支架、反代写约束、知识图谱路由和课堂应用流程方面具有一定探索意义，能够较好回应人工智能赋能基础教育的现实需求，也为后续真实课堂试用、学习数据采集和系统迭代奠定了基础。项目材料较为完整，研究过程规范，成果呈现充分。总体来看，项目完成度较高，成果形式较丰富，达到预期结题要求，同意结题。"
)


def set_run_font(run, size: float = 10.5, bold: bool | None = None, color: RGBColor = INK) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def paragraph_style(paragraph, size: float = 10.5, align=None, before: float = 0, after: float = 3, line: float = 1.15) -> None:
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    for run in paragraph.runs:
        set_run_font(run, size=size)


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_paragraph_text(paragraph, text: str, size: float = 10.5, bold: bool = False, align=None) -> None:
    clear_paragraph(paragraph)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def set_cell_margins(cell, top: int = 60, start: int = 70, bottom: int = 60, end: int = 70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:type"), "dxa")
        node.set(qn("w:w"), str(value))


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_layout(table, width_dxa: int, col_widths: list[int] | None = None) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")

    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:type"), "dxa")
    ind.set(qn("w:w"), "0")

    if col_widths:
        grid = table._tbl.tblGrid
        if grid is None:
            grid = OxmlElement("w:tblGrid")
            table._tbl.insert(0, grid)
        for child in list(grid):
            grid.remove(child)
        for width in col_widths:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(width))
            grid.append(grid_col)
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                set_cell_width(cell, col_widths[min(idx, len(col_widths) - 1)])


def set_border(container, edge: str, val: str, size: str = "8", color: str = "000000") -> None:
    pr = container._element.get_or_add_tcPr() if hasattr(container, "_tc") else container._element.get_or_add_trPr()
    # Only table-level borders are handled separately; this helper is kept for cells/rows if needed.


def set_tbl_borders(table, top: bool = True, bottom: bool = True, inside_h: bool = False) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    config = {
        "top": "single" if top else "nil",
        "bottom": "single" if bottom else "nil",
        "left": "nil",
        "right": "nil",
        "insideV": "nil",
        "insideH": "single" if inside_h else "nil",
    }
    for edge, val in config.items():
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), val)
        node.set(qn("w:sz"), "10" if edge in ("top", "bottom") else "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")


def set_row_bottom_border(row, val: str = "single", size: str = "8") -> None:
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ("top", "left", "right"):
            node = borders.find(qn(f"w:{edge}"))
            if node is None:
                node = OxmlElement(f"w:{edge}")
                borders.append(node)
            node.set(qn("w:val"), "nil")
        node = borders.find(qn("w:bottom"))
        if node is None:
            node = OxmlElement("w:bottom")
            borders.append(node)
        node.set(qn("w:val"), val)
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")


def format_three_line_table(table, table_idx: int) -> None:
    prompt_widths = {
        7: [430, 1220, 1900, 1500, 1500, 1930],
        8: [1200, 1900, 1350, 2100, 1930],
        9: [1500, 1750, 1450, 1900, 1880],
    }
    if table_idx in prompt_widths:
        widths = prompt_widths[table_idx]
        font_size = 7.2 if table_idx == 7 else 7.5
        width_dxa = sum(widths)
        margins = (35, 45, 35, 45)
    else:
        widths = None
        font_size = 8.5
        width_dxa = 9000
        margins = (60, 70, 60, 70)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_layout(table, width_dxa, widths)
    set_tbl_borders(table)

    for r_idx, row in enumerate(table.rows):
        if r_idx == 0:
            set_row_bottom_border(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, *margins)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(run, size=font_size, bold=(r_idx == 0))


def paragraph_has_picture(paragraph) -> bool:
    xml = paragraph._p.xml
    return "<w:drawing" in xml or "<w:pict" in xml or "<v:shape" in xml


def resize_picture_paragraphs(cell) -> None:
    max_w = 430
    for paragraph in cell.paragraphs:
        if not paragraph_has_picture(paragraph):
            continue
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for shape in paragraph._p.xpath(".//*[local-name()='shape']"):
            style = shape.get("style") or ""
            width_match = re.search(r"width:([0-9.]+)pt", style)
            height_match = re.search(r"height:([0-9.]+)pt", style)
            if not width_match or not height_match:
                continue
            width = float(width_match.group(1))
            height = float(height_match.group(1))
            if width > max_w:
                ratio = max_w / width
                new_style = re.sub(r"width:[0-9.]+pt", f"width:{max_w:.2f}pt", style)
                new_style = re.sub(r"height:[0-9.]+pt", f"height:{height * ratio:.2f}pt", new_style)
                shape.set("style", new_style)


def is_table_caption(text: str) -> bool:
    return bool(re.match(r"^表\s*\d+", text.strip()))


def is_figure_caption(text: str) -> bool:
    return bool(re.match(r"^图\s*\d*[\s\u3000]*", text.strip())) and not text.strip().startswith("图谱")


def retitle_captions(report_cell) -> None:
    table_no = 1
    figure_no = 1
    for paragraph in report_cell.paragraphs:
        text = " ".join(paragraph.text.split())
        if is_table_caption(text) and table_no <= len(TABLE_TITLES):
            set_paragraph_text(paragraph, f"表{table_no} {TABLE_TITLES[table_no - 1]}", size=9.2, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
            table_no += 1
        elif is_figure_caption(text) and figure_no <= len(FIGURE_TITLES):
            set_paragraph_text(paragraph, f"图{figure_no} {FIGURE_TITLES[figure_no - 1]}", size=9.2, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
            figure_no += 1


def insert_missing_table_captions(report_cell) -> None:
    # Captions are inserted immediately before the uncaptitled nested tables in the report section.
    missing = {
        6: "表7 角色权限与数据流向",
        7: "表8 CT-GraphMAS三层多智能体提示词设计",
        8: "表9 ADAPT伴学模式智能体设计",
        9: "表10 C-O-D-E-R课堂范式智能体设计",
    }
    tbl_index = 0
    for child in list(report_cell._tc.iterchildren()):
        if child.tag.split("}")[-1] != "tbl":
            continue
        if tbl_index in missing:
            p = OxmlElement("w:p")
            child.addprevious(p)
            from docx.text.paragraph import Paragraph

            paragraph = Paragraph(p, report_cell)
            set_paragraph_text(paragraph, missing[tbl_index], size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
        tbl_index += 1


def update_late_table_caption(report_cell) -> None:
    # The expense table originally appears as 表9; after adding prompt-design tables it becomes 表11.
    for paragraph in report_cell.paragraphs:
        text = " ".join(paragraph.text.split())
        if text.startswith("表9 项目经费使用与报销汇总"):
            set_paragraph_text(paragraph, "表11 项目经费使用与报销汇总", size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_mentor_comment(main_table) -> None:
    comment_cell = main_table.rows[12].cells[0]
    comment_cell.text = ""
    p = comment_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(MENTOR_COMMENT)
    set_run_font(run, size=10.5)

    for _ in range(3):
        blank = comment_cell.add_paragraph()
        blank.paragraph_format.space_after = Pt(6)

    sig = comment_cell.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.paragraph_format.space_after = Pt(0)
    sig_run = sig.add_run("指导教师：              2026年  05月  07日")
    set_run_font(sig_run, size=10.5)


def format_report_text(report_cell) -> None:
    for paragraph in report_cell.paragraphs:
        text = " ".join(paragraph.text.split())
        if not text and not paragraph_has_picture(paragraph):
            paragraph.paragraph_format.space_after = Pt(1)
            continue
        if re.match(r"^\d+\s", text):
            paragraph_style(paragraph, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, after=5, line=1.15)
            for run in paragraph.runs:
                set_run_font(run, size=12, bold=True, color=BLUE)
        elif re.match(r"^\d+\.\d+", text):
            paragraph_style(paragraph, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, after=4, line=1.15)
            for run in paragraph.runs:
                set_run_font(run, size=11, bold=True)
        elif text.startswith(("表", "图")):
            paragraph_style(paragraph, size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER, after=5, line=1.1)
        else:
            paragraph_style(paragraph, size=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3, line=1.18)


def main() -> None:
    doc = Document(DOCX)
    main_table = doc.tables[0]
    report_cell = main_table.rows[10].cells[0]

    insert_missing_table_captions(report_cell)
    retitle_captions(report_cell)
    update_late_table_caption(report_cell)
    format_report_text(report_cell)

    for idx, nested in enumerate(report_cell.tables):
        format_three_line_table(nested, idx)

    resize_picture_paragraphs(report_cell)
    add_mentor_comment(main_table)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
