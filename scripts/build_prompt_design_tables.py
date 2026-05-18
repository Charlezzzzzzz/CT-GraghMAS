"""Build concise prompt-design tables for the CT-GraphMAS final report."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT_PATH = DOCS / "CT-GraphMAS核心提示词与教学模式智能体设计表.docx"

BLUE = RGBColor(31, 77, 120)
MID_BLUE = RGBColor(46, 116, 181)
INK = RGBColor(24, 36, 57)
GRAY = RGBColor(90, 96, 108)
HEADER_FILL = "EAF4FF"
SUBTLE_FILL = "F7FBFF"


def set_font(run, size: float = 10.5, color: RGBColor = INK, bold: bool | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_spacing(paragraph, before: float = 0, after: float = 6, line: float = 1.18) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))


def set_cell_margins(cell, top: int = 80, start: int = 105, bottom: int = 80, end: int = 105) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_borders(table, color: str = "B7C7D9", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def table_width(table, width_in: float = 9.2) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_in * 1440)))


def format_table(table, widths: list[float], font_size: float = 8.7) -> None:
    table_width(table, sum(widths))
    table_borders(table)
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if col_idx < len(widths):
                set_cell_width(cell, widths[col_idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                shade_cell(cell, HEADER_FILL)
            elif col_idx == 0:
                shade_cell(cell, SUBTLE_FILL)
            for paragraph in cell.paragraphs:
                set_spacing(paragraph, after=2, line=1.12)
                if row_idx == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_font(run, size=font_size + (0.2 if row_idx == 0 else 0), bold=(row_idx == 0))


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Title", 18, BLUE, 0, 6),
        ("Heading 1", 13.5, MID_BLUE, 12, 5),
        ("Heading 2", 11.5, BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "CT-GraphMAS 核心提示词与教学模式智能体设计表"
    set_font(header.runs[0], size=8.5, color=GRAY)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer.paragraphs[0]
    footer.text = f"大学生创新创业训练计划项目结题材料 · {date.today().isoformat()}"
    set_font(footer.runs[0], size=8.5, color=GRAY)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_intro(doc: Document) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("CT-GraphMAS 核心提示词与教学模式智能体设计表")
    set_font(title.runs[0], size=18, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("用于大创结题报告中“系统开发成果 / 多智能体工作流设计”部分")
    set_font(p.runs[0], size=10, color=GRAY)

    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.add_run(
        "本表以精简方式概括系统中的三层多智能体提示词、ADAPT 伴学模式和 C-O-D-E-R 课堂范式。"
        "其中，CT-GraphMAS 负责底层诊断、图谱路由和反馈呈现，ADAPT 侧重个性化伴学支架，"
        "C-O-D-E-R 侧重项目式课堂中的任务编排与协同开发。"
    )
    set_font(body.runs[0], size=10.2)
    set_spacing(body, after=8)


def add_core_prompt_table(doc: Document) -> None:
    doc.add_heading("一、CT-GraphMAS 三层多智能体提示词设计表", level=1)
    headers = ["层级", "智能体", "角色定位", "主要输入", "主要输出", "提示词约束"]
    rows = [
        ["L1", "学情与任务整合中枢", "生成诊断前置卷，整合学生问题、图谱学情和路由依据。", "学生输入、目标考点、2-hop 图谱上下文、学习者画像", "结构化诊断 JSON：困境、认知状态、目标考点、路由、风险", "只做诊断与路由；不直接给完整代码；必须依据图谱与画像输出可解析 JSON。"],
        ["L2", "P1 问题分解教研员", "从问题规模和模块边界帮助学生缩小任务范围。", "L1 诊断前置卷、任务类型、目标考点", "拆解路径、最小切入点、下一步聚焦对象", "强调输入-处理-输出与最小可运行步骤；避免替学生完成整题。"],
        ["L2", "P2 抽象建模教研员", "帮助学生把真实情境转化为变量、状态和数据结构。", "L1 诊断、学生描述、已有变量或容器使用情况", "核心状态、数据表示建议、建模支架", "关注现实逻辑到程序状态的映射；不脱离学生当前知识阶段扩展过多概念。"],
        ["L2", "P3 算法设计教研员", "梳理执行顺序、条件分支、循环边界和逻辑闭环。", "L1 诊断、学生代码或步骤、前置知识掌握情况", "流程推演、控制流提示、伪代码级支架", "要求学生先说清步骤再写代码；可给伪代码方向，不给可直接提交的完整程序。"],
        ["L2", "P4 评估与调试教研员", "引导学生通过假设、测试和状态观察定位问题。", "L1 诊断、报错信息、输出结果、测试样例", "调试假设、最小复现、边界样例、变量观察建议", "强调读报错、看状态、做验证；输出局部排查动作，不直接公布修复答案。"],
        ["L3", "T1 同伴型支持智能体", "以低压力同伴语气陪伴新手继续尝试。", "L2 聚合策略、学习者画像、学生原问题", "温和反馈、一个低门槛行动、鼓励性追问", "语气支持但不代写；少术语、低负担，适合破冰期学生。"],
        ["L3", "T2 启发型导师智能体", "将 L2 策略转化为苏格拉底式追问和局部提示。", "学生原问题、L2 聚合策略、相关知识、图谱节点", "100-500 字启发式反馈、关键追问、下一步行动", "只使用 L2 素材；不直接给完整代码；保持追问、局部提示和教学克制。"],
        ["L3", "T3 严格教练型智能体", "用明确边界和检查标准训练学生独立修正。", "L2 策略、任务难度、反代写风险、学生画像", "自查要求、验证任务、提交前检查标准", "语气直接但不打击；拒绝全盘代写，要求学生提交尝试与中间推理。"],
        ["S6", "教学门控与反代写检查器", "检查最终反馈是否越界、代写或超出学生认知阶段。", "L3 草稿、任务属性、风险标记、教师设置", "通过、降级或重写建议", "限制完整答案泄露；将高风险输出降级为提示、伪代码或局部检查。"],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    format_table(table, [0.52, 1.25, 2.05, 1.9, 1.9, 2.05], font_size=8.1)


def add_adapt_table(doc: Document) -> None:
    doc.add_heading("二、ADAPT 伴学模式智能体设计表", level=1)
    headers = ["阶段", "教学含义", "对应智能体", "系统任务", "提示词约束"]
    rows = [
        ["A Assess 诊断卡点", "识别学生当前卡在概念、流程、调试还是情绪信心。", "L1 + P1", "生成诊断前置卷，判断卡点类型和支架强度。", "先诊断再提示；不急于给答案，必要时追问缺失信息。"],
        ["D Design 设计支架", "依据最近发展区选择合适支架。", "P1/P2/P4", "选择问题分解、抽象建模或调试验证支架。", "支架必须小步、可执行，并与图谱前置知识一致。"],
        ["A Act 行动尝试", "让学生完成一个可操作的小步骤。", "L3-T1/T2", "把教研策略转为学生能立即尝试的行动。", "只给下一步，不一次性给完整程序；保留学生主动建构空间。"],
        ["P Practice 练习验证", "通过样例检验学生是否真正理解。", "P4 + L3", "生成正常样例、边界样例和状态观察任务。", "强调验证与解释，不把练习变成答案展示。"],
        ["T Transfer 迁移复盘", "把本次经验迁移到新任务。", "L3-T2/T3 + 思维复盘智能体", "形成迁移问题、反思句和后续任务建议。", "总结计算思维方法，而不是只总结某一道题的代码。"],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    format_table(table, [1.25, 2.15, 1.45, 2.45, 2.25], font_size=8.4)


def add_coder_table(doc: Document) -> None:
    doc.add_heading("三、C-O-D-E-R 课堂范式智能体设计表", level=1)
    headers = ["阶段", "教学含义", "对应智能体", "系统任务", "提示词约束"]
    rows = [
        ["C Contextual Activation 情境激活", "从真实课堂任务中澄清用户、需求、输入、输出和限制。", "需求分析智能体 + P1", "把项目任务转化为可分解的编程问题。", "围绕情境和需求提问，不直接进入代码生成。"],
        ["O Organized Orchestration 组织编排", "教师、学生和智能体共同确定路径、分组与支架强度。", "L1 + 教师调度", "结合画像、任务类型和图谱相关知识选择流程。", "依据诊断结果动态路由，避免无差别调用全部智能体。"],
        ["D Development Collaboration 协同开发", "在学生主导下完成变量、结构、流程和局部实现。", "P2/P3 + 配对编程智能体", "提供建模、伪代码和局部提示，支持学生完成表达。", "只给局部支架和最小示例，不替学生写完整程序。"],
        ["E Error Debugging 循环调试", "围绕报错、样例和变量跟踪进行迭代修正。", "P4 + 代码检查智能体", "定位错误类型，设计测试样例，检查边界条件。", "强调运行证据和自查过程，避免直接粘贴修复版代码。"],
        ["R Reflective Reconstruction 反思重构", "把一次任务沉淀为可迁移的计算思维经验。", "L3 + 思维复盘智能体", "生成复盘问题、重构建议和迁移提示。", "关注方法迁移和表达重构，不停留在结果是否正确。"],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    format_table(table, [1.65, 2.0, 1.7, 2.15, 2.0], font_size=8.25)


def build_doc() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_intro(doc)
    add_core_prompt_table(doc)
    add_adapt_table(doc)
    add_coder_table(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_doc()
