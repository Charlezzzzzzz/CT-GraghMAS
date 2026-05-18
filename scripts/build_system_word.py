"""Build the CT-GraphMAS final-project report DOCX with screenshots."""

from __future__ import annotations

import shutil
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from ct_graphmas.prompt_workflow import (  # noqa: E402
    L1_PROMPT,
    L2_PROMPTS,
    L3_TUTOR_PROMPT,
    RUNTIME_DESIGN,
    WORKFLOW_STEPS,
)


DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
GUIDES = ROOT / "static" / "assets" / "guides"
OUT_PATH = DOCS / "CT-GraphMAS系统开发成果结题报告.docx"

FLOW_PREVIEW = Path("/private/tmp/ct_graphmas_business_flow_preview/CT-GraphMAS业务流程图.pptx.png")
FLOW_IMAGE = ASSETS / "ct_graphmas_business_flow.png"

BLUE = RGBColor(31, 77, 120)
MID_BLUE = RGBColor(46, 116, 181)
INK = RGBColor(22, 34, 52)
GRAY = RGBColor(90, 96, 108)
SOFT_BLUE = "EAF4FF"
PALE_BLUE = "F6FAFF"
TABLE_HEAD = "F2F4F7"
CODE_FILL = "F6F8FA"
YELLOW = "FFF2CC"


def ensure_assets() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    if FLOW_PREVIEW.exists():
        shutil.copyfile(FLOW_PREVIEW, FLOW_IMAGE)

    required = [
        FLOW_IMAGE,
        GUIDES / "student-home.png",
        GUIDES / "student-daily.png",
        GUIDES / "student-teaching.png",
        GUIDES / "student-answer.png",
        GUIDES / "student-graph.png",
        GUIDES / "teacher-workbench.png",
        GUIDES / "teacher-knowledge.png",
        GUIDES / "admin-center.png",
        GUIDES / "admin-graph.png",
        ASSETS / "settings-top.png",
        ASSETS / "settings-workflow.png",
    ]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing report image assets:\n" + "\n".join(missing))


def set_run_font(
    run,
    name: str = "Microsoft YaHei",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 8, line_spacing: float = 1.25) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    style_tokens = [
        ("Title", 22, BLUE, 0, 8),
        ("Subtitle", 12, GRAY, 0, 12),
        ("Heading 1", 16, MID_BLUE, 16, 8),
        ("Heading 2", 13, MID_BLUE, 12, 6),
        ("Heading 3", 12, BLUE, 8, 4),
    ]
    for name, size, color, before, after in style_tokens:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "大学生创新创业训练计划项目结题报告 · CT-GraphMAS 系统开发成果"
    set_run_font(header.runs[0], size=8.5, color=GRAY)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer.paragraphs[0]
    footer.text = "CT-GraphMAS · 知识图谱增强多智能体编程辅导系统"
    set_run_font(footer.runs[0], size=8.5, color=GRAY)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_in: float = 6.5, indent_dxa: int = 120) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_in * 1440)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def table_set_borders(table, color: str = "B7C7D9", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def format_table(table, widths: list[float], header: bool = True, font_size: float = 9.0) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, sum(widths))
    table_set_borders(table)
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if col_idx < len(widths):
                set_cell_width(cell, widths[col_idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, after=2, line_spacing=1.16)
                for run in paragraph.runs:
                    set_run_font(run, size=font_size, color=INK)
            if header and row_idx == 0:
                shade_cell(cell, TABLE_HEAD)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        set_run_font(run, size=font_size + 0.2, color=INK, bold=True)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p)
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=3, after=9, line_spacing=1.1)
    run = p.add_run(text)
    set_run_font(run, size=9, color=GRAY)


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.25) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_callout(doc: Document, title: str, body: str, fill: str = SOFT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 6.5)
    table_set_borders(table, color="9CC2E5", size="8")
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.5)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    shade_cell(cell, fill)

    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=3, line_spacing=1.18)
    run = p.add_run(title)
    set_run_font(run, size=10.5, color=BLUE, bold=True)

    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, after=0, line_spacing=1.2)
    run2 = p2.add_run(body)
    set_run_font(run2, size=9.8, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_prompt_block(doc: Document, title: str, prompt: str) -> None:
    h = doc.add_heading(title, level=3)
    h.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 6.5)
    table_set_borders(table, color="C7D3E2", size="6")
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.5)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    shade_cell(cell, CODE_FILL)

    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line_spacing=1.08)
    lines = prompt.splitlines()
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, name="Microsoft YaHei", size=8.7, color=INK)
        if idx < len(lines) - 1:
            run.add_break()
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(44)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("CT-GraphMAS 系统开发成果结题报告")
    set_run_font(r, size=22, color=BLUE, bold=True)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=18)
    r = subtitle.add_run("知识图谱增强多智能体编程辅导系统 · 大创结题报告系统成果部分")
    set_run_font(r, size=12.5, color=GRAY)

    meta = doc.add_table(rows=6, cols=2)
    data = [
        ("成果名称", "CT-GraphMAS 知识图谱增强多智能体编程辅导系统"),
        ("成果类型", "可运行 Web 原型系统、可交互知识图谱、知识库与多智能体工作流"),
        ("面向对象", "学生端、教师端、管理员端三类用户"),
        ("技术实现", "Python 本地服务、SQLite 数据存储、HTML/CSS/JavaScript 前端、图谱式节点关系与 embedding 检索"),
        ("报告用途", "用于大学生创新创业训练计划项目结题报告中的系统开发成果说明"),
        ("生成日期", date.today().isoformat()),
    ]
    for row, (key, value) in zip(meta.rows, data):
        row.cells[0].text = key
        row.cells[1].text = value
    format_table(meta, [1.35, 5.0], header=False, font_size=9.5)
    for row in meta.rows:
        shade_cell(row.cells[0], TABLE_HEAD)
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)

    add_callout(
        doc,
        "成果定位",
        "本系统以初中 Python 编程学习为主要应用场景，把论文实验中的知识图谱、三层多智能体协同、动态路由、相关知识检索和反代写门控转化为可演示、可扩展、可纳入教学流程的原型系统。报告仅描述当前系统已经实现的功能与设计原理，不使用未实现模块作夸大表述。",
    )


def add_overview(doc: Document) -> None:
    h = doc.add_heading("1. 系统开发成果概述", level=1)
    h.paragraph_format.page_break_before = True
    add_body_paragraph(
        doc,
        "CT-GraphMAS 是面向编程学习场景构建的知识图谱增强多智能体系统。系统围绕学生提出的概念疑惑、代码错误和任务卡点开展诊断，并将学生问题映射到 Python 微观考点、前置知识、学情掌握状态和教学语料片段之中。与直接问答式系统相比，本系统更强调学习过程的可解释性：学生端获得的是启发式支架和下一步行动，教师端能够看到任务分配、知识库、交互记录和学情依据，管理员端则负责维护账号、系统配置、模型接入和图谱数据边界。",
    )
    add_body_paragraph(
        doc,
        "系统开发采用轻量化本地 Web 架构，后端以 Python 提供服务接口并管理用户、文档、知识片段、图谱节点、图谱边和交互记录，前端以淡蓝色界面组织三类角色页面。当前原型已经实现登录、角色化导航、学生日常通用模式与教学模式、教师班级任务管理、知识库上传与 chunking、embedding 检索、Neo4j 风格图谱视图、API 与多智能体设置、三层智能体提示词展示等功能，能够作为结题展示系统，也能够作为后续教学实验与论文扩展的基础平台。",
    )


def add_design_thinking(doc: Document) -> None:
    doc.add_heading("2. 设计目标与总体思路", level=1)
    add_body_paragraph(
        doc,
        "系统的设计目标不是把大模型包装成一个单页聊天窗口，而是把编程学习中的“任务、知识、学情、反馈、教师调控”组织成完整流程。学生登录后只看到与自己学习直接相关的班级、任务、学习画像和对话入口；教师登录后重点处理班级任务、教学材料、知识库与学情复盘；管理员登录后维护图谱、账号、API 网关和模型策略。这样的角色划分使系统更接近真实教学场景，避免学生端暴露不必要的管理概念，也便于教师对课堂任务和学生反馈形成持续控制。",
    )
    add_body_paragraph(
        doc,
        "在智能体设计上，系统采用 L1-L2-L3 的三层结构。L1 层把学生原始输入与知识图谱上下文整合为诊断前置卷，L2 层由问题分解、抽象建模、算法设计、评估调试四类计算思维教研智能体并行或动态调用，L3 层把教研策略转化为学生可理解的启发式话术。该结构使“诊断、教研、表达”三个环节相互分离，既提高了输出可控性，也便于在结题报告中解释系统为何能够贴合计算思维培养目标。",
    )


def add_flow_overview(doc: Document) -> None:
    doc.add_heading("3. 系统业务流程设计", level=1)
    add_body_paragraph(
        doc,
        "CT-GraphMAS 的业务流程采用三泳道结构表示。管理员首先完成系统初始化、角色维护、API 配置和知识图谱配置；教师在此基础上创建班级任务、上传教学材料、查看学生交互记录并开展学情复盘；学生则通过班级任务进入日常通用模式或教学模式，提交问题、代码或报错信息后接收系统反馈。学生端生成的交互记录会回流教师端，教师可以据此更新任务和知识库，形成教学闭环。",
    )
    add_picture(doc, FLOW_IMAGE, "图 1  CT-GraphMAS 系统业务流程图（管理员、学生、教师三泳道）", width=4.7)
    add_body_paragraph(
        doc,
        "从流程机制看，系统并不是把知识图谱作为静态展示模块，而是把图谱放入学生提问后的自动处理链条。学生输入首先被映射到 145 个 Python 微观考点中的目标节点，系统随后抽取该节点的前置知识和学生掌握状态，并将这些上下文交给 L1 和 L2 智能体。教师上传的文档也会被切分为相关知识片段并建立图谱连接，使课堂材料能够参与后续检索与反馈生成。",
    )


def add_role_table(doc: Document) -> None:
    doc.add_heading("4. 角色权限与数据流向", level=1)
    add_body_paragraph(
        doc,
        "三类角色在系统中的权限边界围绕真实教学关系设置。学生端强调完成学习任务和获得支架反馈，教师端强调组织教学与观察学情，管理员端强调维护系统运行条件和图谱配置。不同角色看到的是面向自身工作的表述，而不是以第三者视角介绍功能，这也是当前界面多次修改后形成的一个重要设计结果。",
    )

    table = doc.add_table(rows=1, cols=4)
    headers = ["角色", "主要页面", "核心操作", "数据流向"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    rows = [
        (
            "学生",
            "首页、日常通用模式、教学模式、知识图谱、使用说明",
            "查看班级与默认学号，提交问题、代码或报错，阅读启发式反馈与相关知识。",
            "学生输入进入意图映射、图谱检索和多智能体处理，交互记录回流教师端。",
        ),
        (
            "教师",
            "教师首页、工作台、知识库、学情图谱",
            "创建和分配班级任务，上传教学材料，查看学生交互和知识库图谱。",
            "教师任务进入学生端，上传材料转化为相关知识并参与后续检索。",
        ),
        (
            "管理员",
            "管理员中心、系统设置、知识图谱",
            "维护角色与系统配置，设置 API、多智能体策略、教学门控和图谱数据。",
            "系统配置约束模型调用、提示词工作流和图谱展示范围。",
        ),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = value
    format_table(table, [0.78, 1.55, 2.15, 1.95], font_size=8.7)


def add_feature_screens(doc: Document) -> None:
    doc.add_heading("5. 主要功能实现与截图说明", level=1)
    add_body_paragraph(
        doc,
        "系统主要功能按照学生学习、教师教学、管理员维护和知识资源管理四条线展开。以下截图均来自当前本地原型系统，展示了结题阶段已经实现的页面结构和功能入口。截图说明采用成果报告写法，重点解释各页面在系统流程中的作用，而不是重复描述按钮外观。",
    )

    doc.add_heading("5.1 学生端：从班级任务进入多模式伴学", level=2)
    add_picture(doc, GUIDES / "student-home.png", "图 2  学生首页：默认学号、班级任务、学习画像与入口导航", width=6.25)
    add_body_paragraph(
        doc,
        "学生首页以“我当前要完成什么”为中心组织信息，页面不再呈现全局用户管理或系统任务等学生无需理解的内容，而是显示默认学号、所在班级、教师分配的任务、学习画像和最近学习状态。这样的设计使学生登录后的第一屏能够直接进入学习情境，并为后续教学模式选择提供依据。",
    )
    add_picture(doc, GUIDES / "student-daily.png", "图 3  日常通用模式：标准辅导入口与固定滚动对话框", width=6.25)
    add_body_paragraph(
        doc,
        "日常通用模式固定连接标准辅导流程，适合学生在课后或非结构化情境下提出概念疑惑和代码问题。页面采用固定高度、可上下滚动的输入输出区域，能够容纳较长对话而不破坏整体布局；系统在该模式下仍保留知识图谱检索和相关知识注入，但不会让学生手动选择不必要的教学流程。",
    )
    add_picture(doc, GUIDES / "student-teaching.png", "图 4  教学模式：面向课堂流程的 C-O-D-E-R 与 ADAPT 适配", width=6.25)
    add_body_paragraph(
        doc,
        "教学模式面向教师已经组织好的课堂任务，系统可根据教学环节适配 C-O-D-E-R 或 ADAPT 流程。C-O-D-E-R 更适合项目式编程任务中的情境激活、组织编排、协同开发、循环调试和反思重构，ADAPT 更适合个性化诊断、支架、行动、练习和迁移。由于日常通用模式已经固定为标准辅导，教学模式才保留不同教学流程的切换入口，从而避免功能逻辑混乱。",
    )
    add_picture(doc, GUIDES / "student-answer.png", "图 5  教学伴学回答：L1 诊断、L2 协同、L3 呈现与相关知识", width=6.25)
    add_body_paragraph(
        doc,
        "学生提交问题后，系统会在回答页展示启发式反馈、相关知识、图谱命中节点和多智能体运行过程。反馈不是直接输出完整正确代码，而是通过追问、局部提示、变量状态观察和下一步行动引导学生继续思考。页面中的“相关知识”来自知识库 chunking、embedding 检索和图谱连接，能够让学生看到本次回答所依赖的材料依据。",
    )
    add_picture(doc, GUIDES / "student-graph.png", "图 6  学生图谱视图：检索、拖动与前置相关节点查看", width=6.25)
    add_body_paragraph(
        doc,
        "学生端图谱视图用于把抽象知识点转化为可观察的学习路径。学生可以检索目标考点，拖动节点观察局部结构，点击节点查看前置相关节点与掌握信息。该功能将“先修知识缺口”具体化，避免系统只给出笼统建议，也便于学生理解为什么某个报错或卡点与此前知识有关。",
    )

    doc.add_heading("5.2 教师端：班级任务、知识库与学情复盘", level=2)
    add_picture(doc, GUIDES / "teacher-workbench.png", "图 7  教师工作台：班级任务分配、课堂交互与任务编排", width=6.25)
    add_body_paragraph(
        doc,
        "教师工作台承担课堂组织功能。教师可以创建任务、分配班级、查看任务状态和最近交互，并据此判断学生当前是停留在概念理解、算法组织还是调试验证阶段。学生端看到的任务来自教师端控制，而不是学生自行接触系统管理任务，这使系统更符合课堂教学中的权责关系。",
    )
    add_picture(doc, GUIDES / "teacher-knowledge.png", "图 8  知识库页面：上传文件、生成相关知识、embedding 与图谱视图", width=6.25)
    add_body_paragraph(
        doc,
        "知识库是教师端的重要扩展模块。教师上传课堂材料、案例文本或任务说明后，系统会将文件解析为多个相关知识片段，并生成向量表示用于后续检索；同时，知识片段会与命中的考点节点建立图谱关系。这样，教学材料不只是被静态保存，而是成为学生对话时可被调用、可被解释、可被追溯的知识资源。",
    )

    doc.add_heading("5.3 管理员端：系统配置、图谱维护与安全边界", level=2)
    add_picture(doc, GUIDES / "admin-center.png", "图 9  管理员中心：系统状态、角色入口与维护概览", width=6.25)
    add_body_paragraph(
        doc,
        "管理员中心用于观察系统运行状态和进入配置页面。管理员并不参与具体学习对话，而是维护系统所依赖的用户、班级、模型、图谱和提示词工作流。该页面体现了系统工程层面的可维护性，使结题展示不只是一个学生端界面，而是包含后台治理能力的完整系统。",
    )
    add_picture(doc, ASSETS / "settings-top.png", "图 10  设置页面上部：学习者画像、多智能体偏好、API 与教学模式配置", width=6.25)
    add_body_paragraph(
        doc,
        "设置页将学习者画像、多智能体偏好、API 网关和教学模式参数集中管理。多智能体偏好支持系统动态分配，能够依据 L1 诊断结果和路由结论选择 P1、P2、P3、P4，而不是简单让用户手动堆叠智能体。API 配置以 Base URL、API Key、对话模型和 embedding 模型为核心，原型说明中明确不把密钥写入源码，从而符合系统演示与后续部署的安全要求。",
    )
    add_picture(doc, ASSETS / "settings-workflow.png", "图 11  设置页面下部：核心提示词、工作流步骤与运行时设计", width=6.25)
    add_body_paragraph(
        doc,
        "设置页面下部展示系统工作流和核心提示词，使研究者能够检查 L1、L2、L3 的职责边界。该设计把实验脚本中的提示词与跑批逻辑转化为可读、可配置的系统元数据，便于后续在论文实验、课堂应用或其他教学流程中继续调整。",
    )
    add_picture(doc, GUIDES / "admin-graph.png", "图 12  管理员图谱视图：145 个 Python 微观考点、学情节点与依赖关系", width=6.25)
    add_body_paragraph(
        doc,
        "管理员图谱视图展示了 145 个 Python 微观考点、认知阶段、学生学情状态和前置依赖关系。系统当前使用中等生画像作为演示学情，节点详情能够显示考点编号、认知阶段、掌握度、近期错误次数和前置相关节点。这一功能把论文中的微观知识结构和学情建模落实到可交互界面中，也为教师复盘和后续 Neo4j 迁移提供了结构基础。",
    )


def add_kg_and_kb_design(doc: Document) -> None:
    doc.add_heading("6. 知识图谱与知识库设计原理", level=1)
    add_body_paragraph(
        doc,
        "系统知识图谱由概念节点、学生节点、文档节点、相关知识片段节点和关系边共同构成。概念层面，系统纳入 145 个带认知阶段标记的 Python 微观考点，覆盖执行环境、变量与数据类型、表达式、字符串、分支、循环、列表、字典、模块、函数、算法思想、异常诊断和扩展认知等内容；关系层面，系统通过 REQUIRES 表示前置依赖，通过 KNOWS 表示学生对考点的掌握度、近期错误次数和状态标签，通过 MENTIONS 表示文档片段与知识点之间的命中关系。",
    )
    add_body_paragraph(
        doc,
        "知识库模块承担图谱外部语料进入系统的通道。教师上传文件后，系统对文本进行分块处理，将较长材料转化为适合检索和注入对话的相关知识片段，并为片段生成 embedding。对话发生时，系统不会把整篇文档直接塞入模型上下文，而是根据学生输入、目标考点和图谱邻接关系选取少量高相关片段。这一设计与论文中的图谱增强结论一致，即对初中编程学习而言，局部、可解释、围绕前置知识的增强通常比大规模无差别注入更有利于降低认知负荷。",
    )


def add_prompt_design(doc: Document) -> None:
    doc.add_heading("7. 多智能体提示词与工作流设计", level=1)
    add_body_paragraph(
        doc,
        "CT-GraphMAS 的提示词设计采用“职责分离”原则。L1 不负责最终解答，而是把学生输入、知识图谱命中和 2-hop 学情上下文整合成结构化诊断；L2 不直接面向学生输出，而是从计算思维维度生成教研支架；L3 才负责把 L2 聚合策略改写为学生可读的启发式反馈。这样可以减少单一模型在同一轮对话中同时承担诊断、推理、表达和安全控制所带来的漂移风险。",
    )
    add_body_paragraph(
        doc,
        "L2 层包含 P1 问题分解、P2 抽象建模、P3 算法设计和 P4 评估调试四类智能体。系统默认支持动态路由：当 L1 判断学生卡在问题规模和模块边界时，优先调用 P1；当学生难以把现实情境映射为变量、列表或字典时，调用 P2；当问题集中在步骤顺序、条件流转或循环终止时，调用 P3；当问题表现为报错、输出异常或调试困难时，调用 P4。实验扩展时，系统仍可复现 P1-P4 的多组合比较，但产品原型不把“多智能体数量越多”作为默认目标。",
    )

    table = doc.add_table(rows=1, cols=4)
    headers = ["层级", "角色定位", "输入与输出", "提示词约束"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    rows = [
        ("L1", "学情与任务整合中枢", "输入学生问题、K 节点映射和 2-hop 学情；输出诊断前置卷 JSON。", "不输出代码或最终解法，只界定情境、认知状态、目标考点、路由和风险。"),
        ("P1", "问题分解教研员", "输入 L1 诊断前置卷；输出最小切入点和拆解步骤。", "缩小问题范围，建立最小可运行任务。"),
        ("P2", "抽象建模教研员", "输入 L1 诊断前置卷；输出数据状态和建模支架。", "强调变量、数据容器和现实逻辑的映射。"),
        ("P3", "算法设计教研员", "输入 L1 诊断前置卷；输出执行顺序和控制流支架。", "关注分支、循环、终止条件和逻辑闭环。"),
        ("P4", "评估与调试教研员", "输入 L1 诊断前置卷；输出假设、测试和状态追踪策略。", "通过最小复现、边界样例和变量观察定位问题。"),
        ("L3", "T2 启发型导师", "输入学生原问题和 L2 聚合策略；输出 100-500 字反馈。", "只使用 L2 素材，不直接给完整代码，采用追问和局部提示。"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = value
    format_table(table, [0.55, 1.35, 2.3, 2.15], font_size=8.5)

    add_body_paragraph(
        doc,
        "运行时工作流从 API 网关配置开始，依次经过意图映射、2-hop 图谱上下文抽取、L1 诊断前置卷、L2 计算思维教研、L3 启发呈现和教学门控。API Key 由设置页配置并以掩码显示，源码中不保存密钥；无外部 API 时系统可使用本地规则智能体演示，启用外部模型后则可接入 OpenAI-compatible 网关和 embedding 模型。",
    )

    workflow = doc.add_table(rows=1, cols=3)
    for cell, value in zip(workflow.rows[0].cells, ["步骤", "名称", "系统含义"]):
        cell.text = value
    for step in WORKFLOW_STEPS:
        cells = workflow.add_row().cells
        cells[0].text = step["id"]
        cells[1].text = step["label"]
        cells[2].text = step["description"]
    format_table(workflow, [0.65, 1.55, 4.15], font_size=8.5)

    runtime_text = (
        f"API 隔离方面，{RUNTIME_DESIGN['api_isolation']} 模型模式方面，{RUNTIME_DESIGN['model_mode']} "
        f"图谱模式方面，{RUNTIME_DESIGN['kg_mode']} 路由模式方面，{RUNTIME_DESIGN['routing_mode']} "
        f"模式绑定方面，{RUNTIME_DESIGN['mode_binding']}"
    )
    add_callout(doc, "运行时设计说明", runtime_text, fill=YELLOW)


def add_full_prompts(doc: Document) -> None:
    h = doc.add_heading("8. 核心提示词完整清单", level=1)
    h.paragraph_format.page_break_before = True
    add_body_paragraph(
        doc,
        "本节完整列出当前系统写入提示词工作流的核心模板，便于在结题报告中说明系统输出并非随意生成，而是受明确角色、输入、输出格式和教学边界约束。以下内容为系统原型当前使用的提示词文本，未包含任何 API Key 或个人密钥信息。",
    )

    add_prompt_block(doc, "8.1 L1 学情与任务整合中枢提示词", L1_PROMPT)
    add_prompt_block(doc, "8.2 L2-P1 问题分解教研员提示词", L2_PROMPTS["P1"])
    add_prompt_block(doc, "8.3 L2-P2 抽象建模教研员提示词", L2_PROMPTS["P2"])
    add_prompt_block(doc, "8.4 L2-P3 算法设计教研员提示词", L2_PROMPTS["P3"])
    add_prompt_block(doc, "8.5 L2-P4 评估与调试教研员提示词", L2_PROMPTS["P4"])
    add_prompt_block(doc, "8.6 L3 T2 启发型导师提示词", L3_TUTOR_PROMPT)


def add_validation_and_value(doc: Document) -> None:
    doc.add_heading("9. 系统运行与应用价值", level=1)
    add_body_paragraph(
        doc,
        "从运行结果看，系统已经形成完整闭环：用户可以通过登录进入不同角色界面，学生能够在日常通用模式和教学模式中提交问题并接收启发式反馈，教师能够维护班级任务和知识库，管理员能够查看系统配置和知识图谱。知识库上传后能够形成相关知识片段和图谱连接，图谱页面能够完成搜索、拖动、节点点击和前置相关节点展示，设置页面能够展示 API 配置、多智能体路由、学习者画像和核心提示词工作流。",
    )
    add_body_paragraph(
        doc,
        "该成果的主要价值在于把论文中的理论结论落实为可操作系统。知识图谱负责提供可解释的知识结构和学情上下文，多智能体负责把计算思维拆解为可调度的教研支架，教学门控负责限制直接代写并保留学生思考空间，知识库负责把教师材料转化为可检索的相关知识。后续若接入真实 Neo4j 服务、真实 embedding 模型和学校统一认证，系统可继续扩展为课堂伴学平台、课后辅导平台或编程学习研究数据采集平台。",
    )


def build_doc() -> None:
    ensure_assets()
    doc = Document()
    style_document(doc)
    add_title_page(doc)
    add_overview(doc)
    add_design_thinking(doc)
    add_flow_overview(doc)
    add_role_table(doc)
    add_feature_screens(doc)
    add_kg_and_kb_design(doc)
    add_prompt_design(doc)
    add_full_prompts(doc)
    add_validation_and_value(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_doc()
